#!/usr/bin/env python3
"""
Streamlit Web App for Synthetic Focus Groups - Enhanced Persona UI
This component adds detailed persona profile management to the web app.
"""

import streamlit as st
import pandas as pd
import json
import sys
import os
from typing import Dict, List, Any, Optional

# Add src to path for imports
sys.path.insert(0, os.path.join(os.path.dirname(__file__), 'src'))

# Import persona-related components
from models.persona import Persona

def show_detailed_persona_manager():
    """
    Persona Manager interface for creating, editing, and viewing detailed personas.
    """
    st.header("👤 Detailed Persona Manager")
    st.markdown("Create and manage comprehensive buyer personas for your synthetic focus groups")
    
    # Persona list overview
    show_persona_overview_table()
    st.divider()
    
    tab1, tab2, tab3, tab4 = st.tabs(["Create New Persona", "View & Edit Personas", "Import/Export", "Web Research Builder"])
    
    with tab1:
        show_persona_creator()
    
    with tab2:
        show_persona_editor()
    
    with tab3:
        show_persona_import_export()
    
    with tab4:
        show_web_research_builder()

def show_persona_creator():
    """Interface for creating a new detailed persona from scratch."""
    st.subheader("🆕 Create New Detailed Persona")
    
    # Persona data dictionary to collect all fields
    persona_data = {}
    
    # Create expandable sections for each category of persona details
    with st.expander("1️⃣ Buyer Avatar Basics", expanded=True):
        col1, col2 = st.columns(2)
        
        with col1:
            persona_data["name"] = st.text_input("Full Name", key="new_name", 
                                               help="The persona's full name")
            persona_data["age"] = st.number_input("Age", min_value=18, max_value=80, value=35, key="new_age")
            persona_data["gender"] = st.selectbox("Gender", 
                                                ["Male", "Female", "Non-binary", "Prefer not to specify"], 
                                                key="new_gender")
            persona_data["education"] = st.text_input("Education", key="new_education",
                                                   help="Highest level of education and relevant certifications")
        
        with col2:
            persona_data["relationship_family"] = st.text_area("Relationship & Family Status", key="new_family",
                                                          height=80, help="Marital status, children, family dynamics")
            persona_data["occupation"] = st.text_input("Occupation", key="new_occupation",
                                                    help="Current job title and type of company")
            persona_data["annual_income"] = st.text_input("Annual Income", key="new_income",
                                                       help="Income range or specific amount (include goals if relevant)")
            persona_data["location"] = st.text_input("Location", key="new_location",
                                                  help="City, state/province, country")
    
    with st.expander("2️⃣ Psychographics & Lifestyle", expanded=False):
        col1, col2 = st.columns(2)
        
        with col1:
            hobbies = st.text_area("Hobbies & Interests", key="new_hobbies", height=100,
                                help="List multiple hobbies separated by commas")
            persona_data["hobbies"] = [hobby.strip() for hobby in hobbies.split(",")] if hobbies else []
            
            community = st.text_area("Community Involvement", key="new_community", height=100,
                                   help="Groups, associations, volunteer work")
            persona_data["community_involvement"] = [item.strip() for item in community.split(",")] if community else []
        
        with col2:
            traits = st.text_area("Personality Traits", key="new_traits", height=100,
                               help="Key personality characteristics, separated by commas")
            persona_data["personality_traits"] = [trait.strip() for trait in traits.split(",")] if traits else []
            
            values = st.text_area("Core Values", key="new_values", height=100,
                               help="Guiding principles and beliefs, separated by commas")
            persona_data["values"] = [value.strip() for value in values.split(",")] if values else []
        
        persona_data["free_time_activities"] = st.text_area("Free Time Activities", key="new_free_time", height=80,
                                                       help="How they spend their time outside of work")
        persona_data["lifestyle_description"] = st.text_area("Lifestyle Description", key="new_lifestyle", height=100,
                                                        help="General overview of their day-to-day lifestyle")
    
    with st.expander("3️⃣ Pains & Challenges", expanded=False):
        col1, col2 = st.columns(2)
        
        with col1:
            struggles = st.text_area("Major Struggles", key="new_struggles", height=150,
                                  help="List the biggest challenges they face, one per line")
            persona_data["major_struggles"] = [struggle.strip() for struggle in struggles.split("\n") if struggle.strip()]
        
        with col2:
            obstacles = st.text_area("Obstacles", key="new_obstacles", height=150,
                                  help="List specific barriers they encounter, one per line")
            persona_data["obstacles"] = [obstacle.strip() for obstacle in obstacles.split("\n") if obstacle.strip()]
        
        persona_data["why_problems_exist"] = st.text_area("Why These Problems Exist", key="new_why_problems", height=100,
                                                     help="Root causes of their challenges and struggles")
    
    with st.expander("4️⃣ Fears & Relationship Impact", expanded=False):
        col1, col2 = st.columns(2)
        
        with col1:
            business_fears = st.text_area("Deep Business Fears", key="new_biz_fears", height=150,
                                       help="List their worst business-related fears, one per line")
            persona_data["deep_fears_business"] = [fear.strip() for fear in business_fears.split("\n") if fear.strip()]
            
            personal_fears = st.text_area("Deep Personal Fears", key="new_personal_fears", height=150,
                                       help="List their worst personal fears, one per line")
            persona_data["deep_fears_personal"] = [fear.strip() for fear in personal_fears.split("\n") if fear.strip()]
        
        with col2:
            persona_data["fear_impact_spouse"] = st.text_area("Impact on Spouse/Partner", key="new_spouse_impact", height=80)
            persona_data["fear_impact_kids"] = st.text_area("Impact on Children", key="new_kids_impact", height=80)
            persona_data["fear_impact_employees"] = st.text_area("Impact on Employees/Team", key="new_employee_impact", height=80)
            persona_data["fear_impact_clients"] = st.text_area("Impact on Clients/Customers", key="new_client_impact", height=80)
            
        remarks = st.text_area("Potential Remarks from Others", key="new_remarks", height=100,
                            help="What others might say about them, one per line")
        persona_data["potential_remarks_from_others"] = [remark.strip() for remark in remarks.split("\n") if remark.strip()]
    
    with st.expander("5️⃣ Previous Attempts & Frustrations", expanded=False):
        col1, col2 = st.columns(2)
        
        with col1:
            agencies = st.text_area("Previous Agencies/Consultants Tried", key="new_agencies", height=80,
                                 help="List separated by commas")
            persona_data["previous_agencies_tried"] = [agency.strip() for agency in agencies.split(",")] if agencies else []
            
            software = st.text_area("Previous Software/Tools Tried", key="new_software", height=80,
                                 help="List separated by commas")
            persona_data["previous_software_tried"] = [tool.strip() for tool in software.split(",")] if software else []
            
            diy = st.text_area("DIY Approaches Tried", key="new_diy", height=80,
                            help="List separated by commas")
            persona_data["diy_approaches_tried"] = [approach.strip() for approach in diy.split(",")] if diy else []
        
        with col2:
            persona_data["why_agencies_failed"] = st.text_area("Why Agencies/Consultants Failed", key="new_why_agencies", height=80)
            persona_data["why_software_failed"] = st.text_area("Why Software/Tools Failed", key="new_why_software", height=80)
            persona_data["why_diy_failed"] = st.text_area("Why DIY Approaches Failed", key="new_why_diy", height=80)
    
    with st.expander("6️⃣ Desired Outcomes", expanded=False):
        col1, col2 = st.columns(2)
        
        with col1:
            biz_results = st.text_area("Tangible Business Results Wanted", key="new_biz_results", height=150,
                                    help="List specific measurable outcomes, one per line")
            persona_data["tangible_business_results"] = [result.strip() for result in biz_results.split("\n") if result.strip()]
            
            personal_results = st.text_area("Tangible Personal Results Wanted", key="new_personal_results", height=150,
                                         help="List specific personal outcomes, one per line")
            persona_data["tangible_personal_results"] = [result.strip() for result in personal_results.split("\n") if result.strip()]
        
        with col2:
            emotional = st.text_area("Emotional Transformations Desired", key="new_emotional", height=150,
                                  help="How they want to feel differently, one per line")
            persona_data["emotional_transformations"] = [trans.strip() for trans in emotional.split("\n") if trans.strip()]
            
            soundbites = st.text_area("'If Only' Soundbites", key="new_soundbites", height=150,
                                   help="Statements starting with 'If only I could...', one per line")
            persona_data["if_only_soundbites"] = [bite.strip() for bite in soundbites.split("\n") if bite.strip()]
    
    with st.expander("7️⃣ Hopes & Dreams", expanded=False):
        col1, col2 = st.columns(2)
        
        with col1:
            recognition = st.text_area("Professional Recognition Goals", key="new_recognition", height=100,
                                    help="List desired accolades and recognition, one per line")
            persona_data["professional_recognition_goals"] = [goal.strip() for goal in recognition.split("\n") if goal.strip()]
            
            financial = st.text_area("Financial Freedom Goals", key="new_financial", height=100,
                                  help="List financial objectives, one per line")
            persona_data["financial_freedom_goals"] = [goal.strip() for goal in financial.split("\n") if goal.strip()]
        
        with col2:
            lifestyle = st.text_area("Lifestyle Upgrade Goals", key="new_lifestyle_goals", height=100,
                                  help="List desired lifestyle improvements, one per line")
            persona_data["lifestyle_upgrade_goals"] = [goal.strip() for goal in lifestyle.split("\n") if goal.strip()]
            
            legacy = st.text_area("Family/Legacy Goals", key="new_legacy", height=100,
                               help="List long-term legacy objectives, one per line")
            persona_data["family_legacy_goals"] = [goal.strip() for goal in legacy.split("\n") if goal.strip()]
        
        persona_data["big_picture_aspirations"] = st.text_area("Big Picture Aspirations", key="new_big_picture", height=100,
                                                          help="Overall life and career aspirations")
    
    with st.expander("8️⃣ How They Want to Be Seen by Others", expanded=False):
        reputation = st.text_area("Desired Reputation", key="new_reputation", height=100,
                               help="How they want to be perceived by others, one per line")
        persona_data["desired_reputation"] = [rep.strip() for rep in reputation.split("\n") if rep.strip()]
        
        statements = st.text_area("Success Statements from Others", key="new_statements", height=100,
                               help="What they want others to say about them, one per line")
        persona_data["success_statements_from_others"] = [statement.strip() for statement in statements.split("\n") if statement.strip()]
    
    with st.expander("9️⃣ Unwanted Outcomes", expanded=False):
        avoid = st.text_area("Things to Avoid", key="new_avoid", height=100,
                          help="What they specifically want to avoid, one per line")
        persona_data["things_to_avoid"] = [thing.strip() for thing in avoid.split("\n") if thing.strip()]
        
        unwanted = st.text_area("Unwanted Quotes", key="new_unwanted", height=100,
                             help="Things they never want to hear, one per line")
        persona_data["unwanted_quotes"] = [quote.strip() for quote in unwanted.split("\n") if quote.strip()]
    
    with st.expander("🔟 Summary", expanded=False):
        persona_data["persona_summary"] = st.text_area("One-Paragraph Persona Summary", key="new_summary", height=150,
                                                   help="Concise overview of the entire persona in one paragraph")
    
    with st.expander("1️⃣1️⃣ Day-in-the-Life Scenario", expanded=False):
        persona_data["ideal_day_scenario"] = st.text_area("Ideal Day Scenario", key="new_ideal_day", height=300,
                                                      help="Detailed hour-by-hour description of their ideal day")
    
    # Submit button
    if st.button("💾 Save New Persona", type="primary"):
        if persona_data["name"]:
            save_persona(persona_data)
            st.success(f"✅ Persona '{persona_data['name']}' created successfully!")
        else:
            st.error("❌ Persona name is required!")

def show_persona_editor():
    """Interface for viewing and editing existing personas."""
    st.subheader("✏️ View & Edit Personas")
    
    # Load existing personas
    personas = load_personas()
    
    if not personas:
        st.info("No saved personas found. Create a new persona first.")
        return
    
    # Select persona to view/edit
    persona_names = [p.name for p in personas]
    selected_name = st.selectbox("Select a persona to view or edit:", persona_names)
    
    selected_persona = next((p for p in personas if p.name == selected_name), None)
    if not selected_persona:
        st.error("Could not find selected persona.")
        return
    
    # Display persona card
    st.divider()
    col1, col2 = st.columns([1, 2])
    
    with col1:
        st.subheader(selected_persona.name)
        st.caption(f"{selected_persona.age} • {selected_persona.gender}")
        st.write(f"📍 {selected_persona.location}")
        st.write(f"💼 {selected_persona.occupation}")
        st.write(f"💰 {selected_persona.annual_income}")
        
        # Action buttons
        edit_mode = st.button("✏️ Edit This Persona")
        delete_mode = st.button("🗑️ Delete This Persona")
        
    with col2:
        st.markdown("### Persona Overview")
        st.write(selected_persona.persona_summary)
        
        # Export options
        with st.expander("Export Options", expanded=False):
            section_options = [
                ("Identity/Provenance", "identity"),
                ("Summary", "summary"),
                ("Personality Traits", "traits"),
                ("Values", "values"),
                ("Community Involvement", "community"),
                ("Major Struggles", "struggles"),
                ("Deep Fears (Business)", "fears_business"),
                ("Previous Software Tried", "prev_software"),
                ("Desired Results (Business)", "results_business"),
                ("Desired Results (Personal)", "results_personal"),
                ("Emotional Transformations", "emotional"),
                ("If Only Soundbites", "if_only"),
                ("Desired Reputation", "reputation"),
                ("Things To Avoid", "avoid"),
                ("Unwanted Quotes", "unwanted"),
                ("Big Picture Aspirations", "big_picture"),
                ("Day-in-the-Life", "day_in_life"),
                ("Sources & Citations", "citations"),
            ]
            default_keys = {k for _, k in section_options}
            selected = st.multiselect(
                "Include sections",
                options=[k for _, k in section_options],
                default=list(default_keys),
                format_func=lambda key: next(lbl for lbl,k in section_options if k==key)
            )
        
        # Export buttons (DOCX / PDF)
        docx_bytes = None
        pdf_bytes = None
        try:
            docx_bytes = export_persona_docx(selected_persona, sections=set(selected))
            st.download_button(
                "⬇️ Download DOCX", docx_bytes,
                file_name=f"{selected_persona.name.replace(' ', '_').lower()}_profile.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                key="dl_docx"
            )
        except Exception as e:
            st.caption(f"DOCX export unavailable: {e}")
        try:
            pdf_bytes = export_persona_pdf(selected_persona, sections=set(selected))
            st.download_button(
                "⬇️ Download PDF", pdf_bytes,
                file_name=f"{selected_persona.name.replace(' ', '_').lower()}_profile.pdf",
                mime="application/pdf",
                key="dl_pdf"
            )
        except Exception as e:
            st.caption(f"PDF export unavailable: {e}")
        
        st.markdown("### Key Struggles")
        for struggle in selected_persona.major_struggles[:3]:
            st.write(f"• {struggle}")
        
        st.markdown("### Signature Quote")
        if selected_persona.if_only_soundbites:
            st.info(f"*\"{selected_persona.if_only_soundbites[0]}\"*")
    
    # Display full persona details in expandable sections
    st.divider()
    
    with st.expander("View Complete Persona Profile", expanded=False):
        tab1, tab2, tab3, tab4 = st.tabs(["Basic & Lifestyle", "Challenges & Fears", "Goals & Aspirations", "Life Scenario"])
        
        with tab1:
            col1, col2 = st.columns(2)
            
            with col1:
                st.markdown("#### Basic Information")
                st.write(f"**Education:** {selected_persona.education}")
                st.write(f"**Family:** {selected_persona.relationship_family}")
                
                st.markdown("#### Personality")
                st.write("**Traits:**")
                for trait in selected_persona.personality_traits:
                    st.write(f"• {trait}")
                
                st.write("**Values:**")
                for value in selected_persona.values:
                    st.write(f"• {value}")
            
            with col2:
                st.markdown("#### Lifestyle")
                st.write(f"**Free Time:** {selected_persona.free_time_activities}")
                st.write(f"**Description:** {selected_persona.lifestyle_description}")
                
                st.markdown("#### Community & Hobbies")
                st.write("**Hobbies:**")
                for hobby in selected_persona.hobbies:
                    st.write(f"• {hobby}")
                
                st.write("**Community Involvement:**")
                for community in selected_persona.community_involvement:
                    st.write(f"• {community}")
        
        with tab2:
            col1, col2 = st.columns(2)
            
            with col1:
                st.markdown("#### Major Struggles")
                for struggle in selected_persona.major_struggles:
                    st.write(f"• {struggle}")
                
                st.markdown("#### Business Fears")
                for fear in selected_persona.deep_fears_business:
                    st.write(f"• {fear}")
                
                st.markdown("#### Previous Solutions Tried")
                st.write("**Software:**")
                for software in selected_persona.previous_software_tried:
                    st.write(f"• {software}")
                st.write(f"**Why Failed:** {selected_persona.why_software_failed}")
            
            with col2:
                st.markdown("#### Personal Fears")
                for fear in selected_persona.deep_fears_personal:
                    st.write(f"• {fear}")
                
                st.markdown("#### Fear Impact")
                st.write(f"**On Family:** {selected_persona.fear_impact_spouse}")
                st.write(f"**On Children:** {selected_persona.fear_impact_kids}")
                st.write(f"**On Work:** {selected_persona.fear_impact_employees}")
                
                st.markdown("#### What Others Might Say")
                for remark in selected_persona.potential_remarks_from_others:
                    st.write(f"• {remark}")
        
        with tab3:
            col1, col2 = st.columns(2)
            
            with col1:
                st.markdown("#### Desired Business Outcomes")
                for result in selected_persona.tangible_business_results:
                    st.write(f"• {result}")
                
                st.markdown("#### Professional Recognition")
                for goal in selected_persona.professional_recognition_goals:
                    st.write(f"• {goal}")
                
                st.markdown("#### Financial Goals")
                for goal in selected_persona.financial_freedom_goals:
                    st.write(f"• {goal}")
            
            with col2:
                st.markdown("#### Desired Personal Outcomes")
                for result in selected_persona.tangible_personal_results:
                    st.write(f"• {result}")
                
                st.markdown("#### Emotional Transformations")
                for transformation in selected_persona.emotional_transformations:
                    st.write(f"• {transformation}")
                
                st.markdown("#### Big Picture")
                st.write(selected_persona.big_picture_aspirations)
        
        with tab4:
            st.markdown("#### Ideal Day Scenario")
            st.write(selected_persona.ideal_day_scenario)
    
    # Handle edit mode
    if edit_mode:
        st.divider()
        st.subheader(f"Edit {selected_persona.name}")
        # In a real implementation, this would populate a form with the persona's data
        # and allow updates, similar to the creation form but pre-filled
        st.info("Edit functionality would be implemented here with pre-filled form fields")
    
    # Handle delete mode
    if delete_mode:
        st.divider()
        st.warning(f"Are you sure you want to delete {selected_persona.name}?")
        if st.button("✅ Confirm Delete"):
            # In a real implementation, this would remove the persona from storage
            st.error(f"Deleted {selected_persona.name}")

def citation_store_path() -> str:
    import os
    base = os.path.join(os.path.dirname(__file__), 'data', 'citations')
    os.makedirs(base, exist_ok=True)
    return os.path.join(base, 'citations.json')


def load_citation_store() -> Dict[str, Any]:
    import json, os
    path = citation_store_path()
    if os.path.exists(path):
        try:
            with open(path, 'r', encoding='utf-8') as f:
                return json.load(f)
        except Exception:
            return {}
    return {}


def compute_citation_key(q: Dict[str, Any]) -> str:
    import hashlib
    base = (q.get('url','') + '|' + (q.get('span','')[:200] or '')).encode('utf-8', errors='ignore')
    return hashlib.sha1(base).hexdigest()


def persist_citations(quotes: List[Dict[str, Any]]):
    import json, time
    store = load_citation_store()
    modified = False
    for q in quotes or []:
        key = compute_citation_key(q)
        if key not in store:
            store[key] = {
                'community': q.get('community','external'),
                'url': q.get('url',''),
                'span': q.get('span','')[:500],
                'count': 1,
                'first_seen': time.strftime('%Y-%m-%dT%H:%M:%SZ', time.gmtime()),
                'last_seen': time.strftime('%Y-%m-%dT%H:%M:%SZ', time.gmtime()),
            }
            modified = True
        else:
            store[key]['count'] = int(store[key].get('count', 0)) + 1
            store[key]['last_seen'] = time.strftime('%Y-%m-%dT%H:%M:%SZ', time.gmtime())
            modified = True
    if modified:
        with open(citation_store_path(), 'w', encoding='utf-8') as f:
            json.dump(store, f, indent=2, ensure_ascii=False)


def show_web_research_builder():
    """Build personas from web evidence (Reddit/Quora) + human coordinator sources."""
    import sys, os, io
    sys.path.insert(0, os.path.join(os.path.dirname(__file__), 'src'))
    from research.web_persona_builder import WebPersonaBuilder

    st.subheader("🌐 Web Research Builder")
    st.write("Construct evidence-backed personas from Reddit/Quora and attach coordinator-provided sources (URLs, documents).")

    # Session store for extra sources
    if 'web_extra_sources' not in st.session_state:
        st.session_state.web_extra_sources = []

    with st.form("web_research_sources_form"):
        st.markdown("### Add Coordinator Sources")
        url_text = st.text_area("Source URLs (one per line)")
        files = st.file_uploader(
            "Attach documents (PDF, DOCX, CSV, XLSX)",
            type=["pdf", "docx", "csv", "xlsx"],
            accept_multiple_files=True,
            key="web_sources_upload"
        )
        add_sources = st.form_submit_button("➕ Add Sources")

    if add_sources:
        added = 0
        # URLs
        for line in (url_text or "").splitlines():
            url = line.strip()
            if not url:
                continue
            snippet = fetch_url_text(url)
            if snippet:
                st.session_state.web_extra_sources.append({
                    'source_type': 'url', 'community': 'external_url', 'published_at': '',
                    'weight': 1, 'snippet': snippet[:2000], 'url': url
                })
                added += 1
        # Files
        for f in (files or []):
            try:
                text = extract_file_text(f)
                st.session_state.web_extra_sources.append({
                    'source_type': 'file', 'community': f.name, 'published_at': '',
                    'weight': 1, 'snippet': text[:4000], 'url': ''
                })
                added += 1
            except Exception as e:
                st.error(f"Failed to read {f.name}: {e}")
        if added:
            st.success(f"Added {added} sources")
        else:
            st.info("No sources added")

    with st.expander("Current added sources", expanded=False):
        st.write(f"Total sources: {len(st.session_state.web_extra_sources)}")
        for i, s in enumerate(st.session_state.web_extra_sources, 1):
            st.write(f"{i}. {s.get('community','external')} - {s.get('url','')}")
        if st.button("Clear Sources"):
            st.session_state.web_extra_sources = []
            st.success("Cleared")

    with st.form("web_research_form"):
        st.markdown("### Evidence Query (web)")
        query = st.text_input("Search query (topic, pain point, product, etc.)", value="social media management tool")
        subs = st.text_input("Subreddits (comma-separated)", value="smallbusiness, marketing")
        min_upvotes = st.slider("Minimum upvotes", 0, 100, 10)
        limit = st.slider("Quote limit", 5, 50, 15)
        use_quora_stub = st.checkbox("Include Quora stub if Reddit is empty", value=True)
        submitted = st.form_submit_button("🔎 Gather Evidence & Build Persona", type="primary")

    if submitted:
        try:
            builder = WebPersonaBuilder()
            sub_list = [s.strip() for s in subs.split(',') if s.strip()]
            quotes = builder.gather_quotes(query, subreddits=sub_list, min_upvotes=min_upvotes, limit=limit)
            if not quotes and use_quora_stub:
                quotes = builder.gather_quora_quotes(query, limit=limit)

            if not quotes and not st.session_state.web_extra_sources:
                st.warning("No web quotes found and no coordinator sources provided. Add sources or broaden your query.")
                return

            persona_dict = builder.build_persona_from_evidence(
                query, subreddits=sub_list, min_upvotes=min_upvotes, limit=limit,
                additional_sources=st.session_state.web_extra_sources
            )

            st.success("✅ Persona synthesized from evidence + coordinator sources!")
            st.json({k: v for k, v in persona_dict.items() if k != 'evidence_quotes'})

            with st.expander("📜 Evidence & Sources"):
                for q in persona_dict.get('evidence_quotes', []):
                    st.write(f"• [{q.get('community','external')}] {q.get('span','')[:160]}{'...' if len(q.get('span',''))>160 else ''}")
                    if q.get('url'):
                        st.caption(q['url'])

            save_sources = st.checkbox("Save sources/citations with this persona", value=True)
            if st.button("💾 Save Persona to Library"):
                if save_sources:
                    try:
                        persist_citations(persona_dict.get('evidence_quotes', []))
                    except Exception as e:
                        st.caption(f"Note: could not persist citations store: {e}")
                else:
                    persona_dict.pop('evidence_quotes', None)
                save_persona(persona_dict)
                st.success("Saved to local personas store")
        except Exception as e:
            st.error(f"Error building persona: {e}")


def fetch_url_text(url: str) -> str:
    """Fetch and extract main text from a URL (best-effort)."""
    try:
        import requests
        html = requests.get(url, timeout=10, headers={"User-Agent":"PersonaResearchBot/1.0"}).text
        try:
            from bs4 import BeautifulSoup
            soup = BeautifulSoup(html, 'html.parser')
            # remove scripts/styles
            for tag in soup(["script","style","noscript"]):
                tag.extract()
            text = soup.get_text(" ")
        except Exception:
            # Fallback: crude text
            text = html
        return text[:8000]
    except Exception:
        return ""


def extract_file_text(uploaded_file) -> str:
    """Extract text from uploaded documents."""
    import pandas as pd
    name = uploaded_file.name.lower()
    if name.endswith('.pdf'):
        try:
            import PyPDF2
        except ImportError:
            raise Exception("PyPDF2 not installed. pip install PyPDF2")
        reader = PyPDF2.PdfReader(uploaded_file)
        pages = []
        for p in reader.pages:
            try:
                pages.append(p.extract_text() or "")
            except Exception:
                continue
        return "\n".join(pages)
    elif name.endswith('.docx'):
        try:
            import docx
        except ImportError:
            raise Exception("python-docx not installed. pip install python-docx")
        doc = docx.Document(uploaded_file)
        return "\n".join([para.text for para in doc.paragraphs])
    elif name.endswith('.csv'):
        df = pd.read_csv(uploaded_file)
        return df.to_csv(index=False)
    elif name.endswith('.xlsx'):
        try:
            df = pd.read_excel(uploaded_file)
        except Exception:
            raise Exception("openpyxl required for Excel files. pip install openpyxl")
        return df.to_csv(index=False)
    elif name.endswith('.doc'):
        raise Exception(".doc not supported. Convert to .docx or install a converter (e.g., unoconv).")
    else:
        raise Exception("Unsupported file type")

def show_persona_overview_table():
    """Show a sortable table of all personas: Name, profession, age, project."""
    # Load personas from storage/session
    personas = load_personas()
    if not personas:
        st.info("No saved personas yet. Create one below or import.")
        return
    data = []
    for p in personas:
        data.append({
            'Name': p.name,
            'Profession': p.occupation,
            'Age': p.age,
            'Project': p.created_for_project_name or "",
        })
    df = pd.DataFrame(data)
    st.subheader("All Personas")
    st.dataframe(df.sort_values(by=['Name']), use_container_width=True)
    
    # Export controls
    col1, col2, col3 = st.columns(3)
    with col1:
        if st.button("Export CSV"):
            csv_bytes = df.to_csv(index=False).encode('utf-8')
            st.download_button("Download CSV", csv_bytes, file_name="personas.csv", mime="text/csv")
        if st.button("Export All Profiles (DOCX ZIP)"):
            try:
                import io, zipfile
                personas_all = load_personas()
                bio = io.BytesIO()
                with zipfile.ZipFile(bio, 'w', zipfile.ZIP_DEFLATED) as zf:
                    for p in personas_all:
                        try:
                            zf.writestr(f"{p.name.replace(' ','_').lower()}_profile.docx", export_persona_docx(p))
                        except Exception:
                            continue
                st.download_button("Download Profiles DOCX ZIP", bio.getvalue(), file_name="personas_profiles_docx.zip", mime="application/zip")
            except Exception as e:
                st.error(f"ZIP export failed: {e}")
    with col2:
        # Prepare XLSX export (requires openpyxl)
        if st.button("Export XLSX"):
            try:
                import io
                bio = io.BytesIO()
                with pd.ExcelWriter(bio, engine='openpyxl') as writer:
                    df.to_excel(writer, index=False, sheet_name='Personas')
                st.download_button("Download XLSX", bio.getvalue(), file_name="personas.xlsx", mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet")
            except Exception:
                st.error("Install openpyxl to export XLSX: pip install openpyxl")
        if st.button("Export All Profiles (PDF ZIP)"):
            try:
                import io, zipfile
                personas_all = load_personas()
                bio = io.BytesIO()
                with zipfile.ZipFile(bio, 'w', zipfile.ZIP_DEFLATED) as zf:
                    for p in personas_all:
                        try:
                            zf.writestr(f"{p.name.replace(' ','_').lower()}_profile.pdf", export_persona_pdf(p))
                        except Exception:
                            continue
                st.download_button("Download Profiles PDF ZIP", bio.getvalue(), file_name="personas_profiles_pdf.zip", mime="application/zip")
            except Exception as e:
                st.error(f"ZIP export failed: {e}")
    with col3:
        if st.button("Export DOCX"):
            try:
                import docx, io
                doc = docx.Document()
                doc.add_heading('Personas', level=1)
                table = doc.add_table(rows=1, cols=len(df.columns))
                hdr_cells = table.rows[0].cells
                for i, col in enumerate(df.columns):
                    hdr_cells[i].text = str(col)
                for _, row in df.iterrows():
                    cells = table.add_row().cells
                    for i, col in enumerate(df.columns):
                        cells[i].text = str(row[col])
                bio = io.BytesIO()
                doc.save(bio)
                st.download_button("Download DOCX", bio.getvalue(), file_name="personas.docx", mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
            except Exception:
                st.error("Install python-docx to export DOCX: pip install python-docx")


def export_persona_docx(p, sections: set | None = None) -> bytes:
    import io, docx
    doc = docx.Document()
    doc.add_heading(p.name, level=1)
    meta = doc.add_paragraph()
    if not sections or 'identity' in sections:
        meta.add_run(f"Profession: {p.occupation}\n").bold = True
        meta.add_run(f"Age: {p.age}\n")
        if getattr(p, 'education', None):
            meta.add_run(f"Education: {getattr(p, 'education', '')}\n")
    if getattr(p, 'annual_income', None):
        meta.add_run(f"Income: {getattr(p, 'annual_income', '')}\n")
    if p.location:
        meta.add_run(f"Location: {p.location}\n")
    if getattr(p, 'relationship_family', None):
        meta.add_run(f"Family: {getattr(p, 'relationship_family', '')}\n")
    if p.created_for_project_name:
        meta.add_run(f"Project: {p.created_for_project_name}\n")
    
    if not sections or 'summary' in sections:
        doc.add_heading("Summary", level=2)
        doc.add_paragraph(p.persona_summary or "")
    
    def add_list(title, items):
        if items:
            doc.add_heading(title, level=2)
            for it in items:
                doc.add_paragraph(str(it), style='List Bullet')
    
    if not sections or 'traits' in sections:
        add_list("Personality Traits", getattr(p, 'personality_traits', []))
    if not sections or 'values' in sections:
        add_list("Values", getattr(p, 'values', []))
    if not sections or 'community' in sections:
        add_list("Community Involvement", getattr(p, 'community_involvement', []))
    if not sections or 'struggles' in sections:
        add_list("Major Struggles", p.major_struggles)
    if not sections or 'fears_business' in sections:
        add_list("Deep Fears (Business)", p.deep_fears_business)
    if not sections or 'prev_software' in sections:
        add_list("Previous Software Tried", p.previous_software_tried)
    if not sections or 'results_business' in sections:
        add_list("Desired Results (Business)", p.tangible_business_results)
    if not sections or 'results_personal' in sections:
        add_list("Desired Results (Personal)", getattr(p, 'tangible_personal_results', []))
    if not sections or 'emotional' in sections:
        add_list("Emotional Transformations", getattr(p, 'emotional_transformations', []))
    if not sections or 'if_only' in sections:
        add_list("If Only Soundbites", p.if_only_soundbites)
    if not sections or 'reputation' in sections:
        add_list("Desired Reputation", getattr(p, 'desired_reputation', []))
    if not sections or 'unwanted' in sections:
        add_list("Unwanted Quotes", getattr(p, 'unwanted_quotes', []))
    if not sections or 'avoid' in sections:
        add_list("Things To Avoid", getattr(p, 'things_to_avoid', []))
    
    if (not sections or 'big_picture' in sections) and getattr(p, 'big_picture_aspirations', None):
        doc.add_heading("Big Picture Aspirations", level=2)
        doc.add_paragraph(p.big_picture_aspirations)
    if (not sections or 'day_in_life' in sections) and getattr(p, 'ideal_day_scenario', None):
        doc.add_heading("Day-in-the-Life", level=2)
        doc.add_paragraph(p.ideal_day_scenario)

    # Citations
    if (not sections or 'citations' in sections) and getattr(p, 'evidence_quotes', None):
        doc.add_heading("Sources & Citations", level=2)
        for q in p.evidence_quotes[:20]:
            line = f"[{q.get('community','external')}] {q.get('span','')[:180]}"
            doc.add_paragraph(line)
            if q.get('url'):
                doc.add_paragraph(q['url'])
        # Citation index using global store counts
        try:
            store = load_citation_store()
            index = {}
            for q in p.evidence_quotes:
                key = compute_citation_key(q)
                meta = store.get(key)
                if meta:
                    index[key] = meta
            if index:
                doc.add_heading("Citation Index", level=2)
                for key, meta in index.items():
                    doc.add_paragraph(f"[{meta.get('community','external')}] count={meta.get('count',1)}")
                    if meta.get('url'):
                        doc.add_paragraph(meta.get('url'))
                    if meta.get('span'):
                        doc.add_paragraph(meta.get('span')[:180])
        except Exception:
            pass
    
    bio = io.BytesIO()
    doc.save(bio)
    return bio.getvalue()


def export_persona_pdf(p, sections: set | None = None) -> bytes:
    # Simple PDF using reportlab
    import io
    from reportlab.lib.pagesizes import letter
    from reportlab.pdfgen import canvas
    from reportlab.lib.units import inch
    
    def draw_wrapped(c, text, x, y, max_width=7.0*inch, leading=14):
        from reportlab.pdfbase.pdfmetrics import stringWidth
        words = text.split()
        line = ""
        while words and y > 1*inch:
            while words and stringWidth(line + (words[0] + " "), "Helvetica", 11) < max_width:
                line += words.pop(0) + " "
            c.drawString(x, y, line.strip())
            y -= leading
            line = ""
        return y
    
    bio = io.BytesIO()
    c = canvas.Canvas(bio, pagesize=letter)
    width, height = letter
    x, y = 1*inch, height - 1*inch
    c.setFont("Helvetica-Bold", 16)
    c.drawString(x, y, p.name)
    y -= 18
    c.setFont("Helvetica", 11)
    meta_parts = []
    if not sections or 'identity' in sections:
        meta_parts = [f"Profession: {p.occupation}", f"Age: {p.age}"]
        if getattr(p, 'education', None):
            meta_parts.append(f"Education: {getattr(p, 'education', '')}")
    if getattr(p, 'annual_income', None):
        meta_parts.append(f"Income: {getattr(p, 'annual_income', '')}")
    if p.location:
        meta_parts.append(f"Location: {p.location}")
    if getattr(p, 'relationship_family', None):
        meta_parts.append(f"Family: {getattr(p, 'relationship_family', '')}")
    if p.created_for_project_name:
        meta_parts.append(f"Project: {p.created_for_project_name}")
    meta = "  |  ".join(meta_parts)
    y = draw_wrapped(c, meta, x, y)
    y -= 10
    if not sections or 'summary' in sections:
        c.setFont("Helvetica-Bold", 12)
        c.drawString(x, y, "Summary")
        y -= 14
        c.setFont("Helvetica", 11)
        y = draw_wrapped(c, p.persona_summary or "", x, y)
        y -= 10
    def write_list(title, items):
        nonlocal x, y
        if not items:
            return
        c.setFont("Helvetica-Bold", 12)
        c.drawString(x, y, title)
        y -= 14
        c.setFont("Helvetica", 11)
        for it in items:
            y = draw_wrapped(c, f"• {str(it)}", x, y)
        y -= 4
    if not sections or 'traits' in sections:
        write_list("Personality Traits", getattr(p, 'personality_traits', []))
    if not sections or 'values' in sections:
        write_list("Values", getattr(p, 'values', []))
    if not sections or 'community' in sections:
        write_list("Community Involvement", getattr(p, 'community_involvement', []))
    if not sections or 'struggles' in sections:
        write_list("Major Struggles", p.major_struggles)
    if not sections or 'fears_business' in sections:
        write_list("Deep Fears (Business)", p.deep_fears_business)
    if not sections or 'prev_software' in sections:
        write_list("Previous Software Tried", p.previous_software_tried)
    if not sections or 'results_business' in sections:
        write_list("Desired Results (Business)", p.tangible_business_results)
    if not sections or 'results_personal' in sections:
        write_list("Desired Results (Personal)", getattr(p, 'tangible_personal_results', []))
    if not sections or 'emotional' in sections:
        write_list("Emotional Transformations", getattr(p, 'emotional_transformations', []))
    if not sections or 'if_only' in sections:
        write_list("If Only Soundbites", p.if_only_soundbites)
    if not sections or 'reputation' in sections:
        write_list("Desired Reputation", getattr(p, 'desired_reputation', []))
    if not sections or 'avoid' in sections:
        write_list("Things To Avoid", getattr(p, 'things_to_avoid', []))
    if not sections or 'unwanted' in sections:
        write_list("Unwanted Quotes", getattr(p, 'unwanted_quotes', []))
    if (not sections or 'big_picture' in sections) and getattr(p, 'big_picture_aspirations', None):
        c.setFont("Helvetica-Bold", 12)
        c.drawString(x, y, "Big Picture Aspirations")
        y -= 14
        c.setFont("Helvetica", 11)
        y = draw_wrapped(c, p.big_picture_aspirations, x, y)
    if (not sections or 'day_in_life' in sections) and getattr(p, 'ideal_day_scenario', None):
        c.setFont("Helvetica-Bold", 12)
        c.drawString(x, y, "Day-in-the-Life")
        y -= 14
        c.setFont("Helvetica", 11)
        y = draw_wrapped(c, p.ideal_day_scenario, x, y)

    # Citations
    if (not sections or 'citations' in sections) and getattr(p, 'evidence_quotes', None):
        c.setFont("Helvetica-Bold", 12)
        c.drawString(x, y, "Sources & Citations")
        y -= 14
        c.setFont("Helvetica", 11)
        quotes = p.evidence_quotes[:12]
        for q in quotes:
            y = draw_wrapped(c, f"[{q.get('community','external')}] {q.get('span','')[:140]}", x, y)
            if q.get('url'):
                y = draw_wrapped(c, q['url'], x, y)
            y -= 4
        # Citation index counts
        try:
            store = load_citation_store()
            index = {}
            for q in p.evidence_quotes:
                key = compute_citation_key(q)
                meta = store.get(key)
                if meta:
                    index[key] = meta
            if index:
                c.setFont("Helvetica-Bold", 12)
                c.drawString(x, y, "Citation Index")
                y -= 14
                c.setFont("Helvetica", 11)
                for key, meta in list(index.items())[:12]:
                    y = draw_wrapped(c, f"[{meta.get('community','external')}] count={meta.get('count',1)}", x, y)
                    if meta.get('url'):
                        y = draw_wrapped(c, meta.get('url'), x, y)
                    if meta.get('span'):
                        y = draw_wrapped(c, meta.get('span')[:140], x, y)
                    y -= 4
        except Exception:
            pass

    c.showPage()
    c.save()
    return bio.getvalue()


def show_persona_import_export():
    """Interface for importing and exporting detailed personas."""
    st.subheader("📤 Import & Export Personas")
    
    tab1, tab2 = st.tabs(["Import", "Export"])
    
    with tab1:
        st.markdown("### Import Personas")
        st.write("Upload personas from file in various formats:")
        
        import_format = st.radio(
            "Select Import Format:",
            ["CSV (Basic)", "CSV (Detailed)", "JSON", "Python File"],
            horizontal=True
        )
        
        if import_format == "CSV (Basic)":
            st.info("Upload a CSV with basic persona fields (name, age, gender, occupation, etc.)")
            upload_file = st.file_uploader("Upload CSV File", type=["csv"], key="basic_csv_upload")
            
            if upload_file:
                try:
                    df = pd.read_csv(upload_file)
                    st.success(f"CSV file loaded with {len(df)} records")
                    st.dataframe(df.head())
                    
                    if st.button("Import These Personas"):
                        st.success(f"Imported {len(df)} personas!")
                except Exception as e:
                    st.error(f"Error loading CSV: {e}")
        
        elif import_format == "CSV (Detailed)":
            st.info("Upload a CSV with all detailed persona fields")
            upload_file = st.file_uploader("Upload Detailed CSV File", type=["csv"], key="detailed_csv_upload")
            
            # Similar processing as above
            
        elif import_format == "JSON":
            st.info("Upload a JSON file with persona data")
            upload_file = st.file_uploader("Upload JSON File", type=["json"], key="json_upload")
            
            if upload_file:
                try:
                    data = json.load(upload_file)
                    st.success(f"JSON file loaded with {len(data)} personas")
                    st.json(data[0] if data else {})
                    
                    if st.button("Import These Personas"):
                        st.success(f"Imported {len(data)} personas!")
                except Exception as e:
                    st.error(f"Error loading JSON: {e}")
        
        elif import_format == "Python File":
            st.info("Upload a Python file with Persona objects (advanced)")
            upload_file = st.file_uploader("Upload Python File", type=["py"], key="py_upload")
            
            if upload_file:
                content = upload_file.read().decode()
                st.code(content, language="python")
                
                if st.button("Import Python Personas"):
                    st.success("Python persona import would be processed here")
    
    with tab2:
        st.markdown("### Export Personas")
        st.write("Export your personas in various formats:")
        
        # Load personas for export
        personas = load_personas()
        
        if not personas:
            st.info("No personas available for export.")
            return
        
        # Select personas to export
        persona_names = [p.name for p in personas]
        selected_personas = st.multiselect("Select personas to export:", persona_names, default=persona_names)
        
        if not selected_personas:
            st.warning("Please select at least one persona to export.")
            return
        
        # Export format options
        export_format = st.radio(
            "Select Export Format:",
            ["CSV", "JSON", "Python Script"],
            horizontal=True
        )
        
        # Select personas for export
        export_personas = [p for p in personas if p.name in selected_personas]
        
        if export_format == "CSV":
            # Convert personas to DataFrame for CSV export
            if st.button("Generate CSV Export"):
                st.success(f"CSV export ready for {len(export_personas)} personas")
                # In real implementation, generate and provide download
                st.download_button(
                    "Download CSV",
                    "name,age,gender,occupation\nMock Data",
                    "personas_export.csv",
                    "text/csv",
                    key="csv_download"
                )
        
        elif export_format == "JSON":
            # Convert personas to JSON
            if st.button("Generate JSON Export"):
                st.success(f"JSON export ready for {len(export_personas)} personas")
                # In real implementation, convert and provide download
                sample_json = json.dumps([{"name": p.name, "age": p.age} for p in export_personas])
                st.download_button(
                    "Download JSON",
                    sample_json,
                    "personas_export.json",
                    "application/json",
                    key="json_download"
                )
        
        elif export_format == "Python Script":
            # Generate Python code
            if st.button("Generate Python Script"):
                st.success(f"Python script ready for {len(export_personas)} personas")
                # In real implementation, generate actual code
                sample_code = "# Sample persona code\ndef create_personas():\n    # Code would go here\n    pass"
                st.code(sample_code, language="python")
                st.download_button(
                    "Download Python Script",
                    sample_code,
                    "personas_export.py",
                    "text/plain",
                    key="py_download"
                )

def save_persona(persona_data):
    """Save a new persona (mock implementation)."""
    # In a real implementation, this would save to a database or file
    if 'saved_personas' not in st.session_state:
        st.session_state.saved_personas = []
    
    # Convert to Persona object
    try:
        new_persona = Persona(**persona_data)
        st.session_state.saved_personas.append(new_persona)
    except Exception as e:
        st.error(f"Error creating persona: {e}")

def load_personas():
    """Load existing personas (mock implementation)."""
    # In a real implementation, this would load from a database or file
    if 'saved_personas' not in st.session_state:
        # For demo, create sample personas if none exist
        from sample_detailed_personas import (
            create_sarah_marketing_maven,
            create_mike_growth_hacker, 
            create_jenny_scaling_solopreneur
        )
        
        st.session_state.saved_personas = [
            create_sarah_marketing_maven(),
            create_mike_growth_hacker(),
            create_jenny_scaling_solopreneur()
        ]
    
    return st.session_state.saved_personas

if __name__ == "__main__":
    st.set_page_config(
        page_title="Persona Manager",
        page_icon="👤",
        layout="wide"
    )
    show_detailed_persona_manager()